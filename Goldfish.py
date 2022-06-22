import bs4 as bs
import os
import urllib.request
from docx import Document
from docx.shared import Inches
from docx.enum.section import WD_SECTION
from docx.enum.section import WD_ORIENT
from docx.shared import Mm

decknumber = str(2733625)  								#Hardcoded value for debugging
#decknumber = input("Please enter decknumber: ") 		#Ask for decknumber 
prelink = "https://www.mtggoldfish.com/deck/visual/"			#Define first part of url
loginname = os.getlogin()                               #Get Windows user for os_path directory 
link = prelink+decknumber                               #Concat the final link where deck can be found
os_path = 'C:/Users/'+loginname+'/Desktop/Scraping/'    #Build OS path where decks will be downloaded to
decknumber = 'Goldfish_Deck'+decknumber                          #Concat the work Deck and the number of the deck
deck_path = os_path+decknumber+'/'                      #Create deck_path directory variable
deck_txt_file_name = deck_path+str(decknumber)+'.txt'   #Create the deck file name
card_image_path = deck_path+'cards/'					#Create card  directory variable

print('Following Deck will be downloaded: ' + link)     #Show user what link is used for downloading deck
sauce = urllib.request.urlopen(link).read()             #Beautifulsoup object for link reading

if not os.path.exists(deck_path):                       #Check if deck folder allready exists
    os.makedirs(deck_path)                              #If it doesnt exist, create folder where deck is downloaded

if not os.path.exists(card_image_path):                 #Check if cards folder allready exists
    os.makedirs(card_image_path)                        #If it doesnt exist, create folder where cards are downloaded

print('Folder '+decknumber+' created!')					#Show user that folder is created

f = open(deck_txt_file_name, "w")                       #Create deck.txt file
soup = bs.BeautifulSoup(sauce,'lxml')#                  #Load soup object
for aclass in soup.find_all('img'):                     #Only take img tags  from lxml
	f.write(aclass['src'])                              #Write every url to deck.txt file
	f.write('\n')                                       #Puts a break after every line
f.close()                                               #Close

links_to_keep = []
with open(deck_txt_file_name, "r") as f:

     for line in f.readlines():
         if '.jpg' in line:
             links_to_keep.append(line)

# Write all the links in our list to the file
with open(deck_txt_file_name, "w") as f:

    for link in links_to_keep:
        f.write(link)

document = Document()
section = document.sections[0]
section.page_height = Mm(210)
section.page_width = Mm(297)
section.left_margin = Mm(20)
section.right_margin = Mm(20)
section.top_margin = Mm(17)
section.bottom_margin = Mm(17)
section.header_distance = Mm(12.7)
section.footer_distance = Mm(12.7)


p = document.add_paragraph()
r = p.add_run()

response = open(deck_txt_file_name, 'r+')                               		#Read deck txt file
count = 1																		#Set counter
for line in response:                                                   
	urllib.request.urlretrieve(line,card_image_path+'/img'+str(count)+'.jpg') 	#For every url, download the image
	r.add_picture(card_image_path+'/img'+str(count)+'.jpg', width=Inches(2.48031))
	count = count + 1                                                   		#Raise counter for image name
  
print(str(count -1) + ' Cards Downloaded!')                                     # Show user amount of cards downloaded 

document.save('C:/Users/Remsporium/Desktop/Scraping/'+decknumber+'/'+decknumber+'.docx')
print('Printable Word document generated!')	